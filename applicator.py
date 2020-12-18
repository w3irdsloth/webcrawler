 ##############
## Applicator ##
 ##############

import os
from os.path import splitext

class Applicator(object):
    """ Creates an object for applying text to files """
    def __init__(self):
        self.text = ""
        self.tag = ""

    def set_text(self, text):
       self.text = text
    
    def set_tag(self, tag):
        self.tag = tag

    #Apply text to tag in file
    def apply_text(self, document):
        ext = splitext(document)[1]
        if ext == ".txt":
            print("applying text to .txt file...")
            try:
                with open(document, "a") as temp_file:
                    temp_file.write(self.text)
                    temp_file.close()

            except:
                print("text application failed")
                raise SystemExit

        elif ext == ".docx":
            from docx import Document
            print("printing to .docx file...")
            try:
                doc = Document(document)
            
            except:
                doc = Document()

            tag_check = False
            if len(self.tag) >= 1:
                for prgph in doc.paragraphs:
                    if self.tag in prgph.text:
                        print("printing to selected tag...")
                        prgph.text = self.text                    
                        doc.save(document)
                        tag_check = True       
                        break

                if tag_check == False:
                    print("tag not found")
                    print("printing to end of file...")
                    doc.add_paragraph(self.text)
                    doc.save(document)

            else:
                print("no tag selected")
                print("printing to end of file...")
                doc.add_paragraph(self.text)
                doc.save(document)

        else:
            print("unsupported file type")
