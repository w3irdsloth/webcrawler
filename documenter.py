     ################
    ##  Documenter  ##
     ################

import os
from docx import Document

class Documenter(object):
    """Add text to a new or existing document based on a specified placeholder tag """

    def __init__(self):
        self.text = ""
        self.ext = ""
        self.tag = ""

    def get_text(self):
        return self.text

    def set_text(self, text):
        self.text = text

    def set_tag(self, tag):
        self.tag = tag

    #Look for selected document and create new one if it doesn't exist before pasting text to tag or at end of file if tag not found
    def document_text(self, document):
        if os.path.isfile(document):
            try:
                print("printing to .docx file...")
                doc = Document(document)
            
            except:
                print("unsupported document")
                raise SystemExit

        else:
            doc = Document()

        tag_check = False
        if len(self.tag) >= 1:
            for prgph in doc.paragraphs:
                if self.tag in prgph.text:
                    print("printing to selected tag...")
                    prgph.text = self.text                    
                    doc.save(document)
                    tag_check = True       
                    break

                if tag_check == False:
                    print("tag not found")
                    print("printing to end of file...")
                    doc.add_paragraph(self.text)
                    doc.save(document)

        else:
            print("no tag selected")
            print("printing to end of file...")
            doc.add_paragraph(self.text)
            doc.save(document)


