 ##############
## Applicator ##
 ##############

import os
from os.path import splitext, isfile

class Applicator(object):
    """ Apply text to file at tag (if tag is provided/exists) """
    def __init__(self):
        self.tag = ""

    def set_tag(self, tag):
        self.tag = tag

    def get_tag(self):
        return self.tag

    def is_tag(self):
        if len(self.tag) > 0:
            print("tag selected")
            return True

        else:
            print("no tag selected")
            return False

    def tag_len(self):
        return len(self.tag)

    #Apply text to tag in file
    def apply_text(self, text, document):
        print("applying text to " + str(document))
        ext = splitext(document)[1]
        if ".txt" in ext or ".md" in ext:
            if os.path.isfile(document):
                f = open(document, "r")
                doc_text = f.read()

            else:
                doc_text = ""

            f = open(document, "w")
            if self.is_tag():
                if self.tag in doc_text:
                    slice_index_start = doc_text.index(self.tag)
                    slice_index_end = slice_index_start + self.tag_len()
                    new_doc_text = ''.join(doc_text[:slice_index_start] + text + doc_text[slice_index_end:])
                    new_text = new_doc_text

                else:
                    print("tag not found")
                    if len(doc_text) > 0:
                        new_text = doc_text + "\n"
                        new_text = new_text + text

            else:
                if len(doc_text) > 0:
                    new_text = doc_text + "\n"
                    new_text = new_text + text

                else:
                    new_text = text

            f.write(new_text)    
            f.close()
            return True

        elif ".doc" in ext:
            from docx import Document

            if os.path.isfile(document):
                doc = Document(document)
            
            else:
                doc = Document()

            if self.is_tag():
                for prgph in doc.paragraphs:
                    if self.tag in prgph.text:
                        prgph_string = prgph.text
                        slice_index_start = prgph_string.index(self.tag)
                        slice_index_end = slice_index_start + self.tag_len()
                        new_prgph_text = ''.join(prgph_string[:slice_index_start] + text + prgph_string[slice_index_end:])
                        prgph.text = new_prgph_text       
                        break

                    else:
                        print("tag not found")
                        doc.add_paragraph(text)

            else:
                doc.add_paragraph(text)
                
            doc.save(document)
            return True

        else:
            return False